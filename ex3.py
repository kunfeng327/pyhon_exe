import tkinter as tk
from tkinter import ttk, messagebox
import random
from docx import Document
import os

def generate_questions():
    try:
        num = int(entry_num.get())
        grade = entry_grade.get().strip()

        questions = []
        for _ in range(num):
            if grade in ["一", "二", "1", "2"]:
                a = random.randint(0, 20)
                b = random.randint(0, 20)
                op = random.choice(["+", "-"])
                if op == "-" and a < b:
                    a, b = b, a
                q = f"{a}{op}{b}="
            elif grade in ["三", "四", "3", "4"]:
                a = random.randint(10, 100)
                b = random.randint(10, 100)
                op = random.choice(["+", "-", "*", "/"])
                if op == "/":
                    b = random.randint(1, 20)
                    a = b * random.randint(1, 10)
                q = f"{a}{op}{b}="
            elif grade in ["五", "六", "5", "6"]:
                a = random.randint(1, 50)
                b = random.randint(1, 20)
                c = random.randint(1, 20)
                op1 = random.choice(["+", "-", "*", "/"])
                op2 = random.choice(["+", "-"])
                q = f"({a}{op1}{b}){op2}{c}="
            else:
                messagebox.showerror("错误", "请输入正确年级")
                return
            questions.append(q)

        doc = Document()
        doc.add_heading("小学口算题", 0)
        rows = (num + 3) // 4
        table = doc.add_table(rows=rows, cols=4)
        table.style = "Table Grid"

        index = 0
        for i in range(rows):
            for j in range(4):
                if index < len(questions):
                    table.cell(i, j).text = questions[index]
                    index += 1

        doc.save("口算题.docx")
        messagebox.showinfo("成功", f"已生成 {num} 道题，保存为：口算题.docx")
        os.startfile("口算题.docx")
    except:
        messagebox.showerror("错误", "输入无效！")

# 界面
root = tk.Tk()
root.title("小学口算题生成器")
root.geometry("300x180")

tk.Label(root, text="Number:").pack(pady=5)
entry_num = ttk.Entry(root)
entry_num.pack(pady=5)
entry_num.insert(0, "200")

tk.Label(root, text="Grade:").pack(pady=5)
entry_grade = ttk.Entry(root)
entry_grade.pack(pady=5)

ttk.Button(root, text="GO", command=generate_questions).pack(pady=10)

root.mainloop()